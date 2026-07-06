"""
Word Document Exporter Service
Export story chapters to Microsoft Word format (.docx)
"""
import os
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from loguru import logger


class WordExporter:
    """Service for exporting stories to Word documents"""

    def __init__(self, output_dir: str = "storage/exports"):
        """
        Initialize exporter

        Args:
            output_dir: Directory to save exported files
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_document(
        self,
        title: str,
        chapters: List[dict],
        author: Optional[str] = None
    ) -> Document:
        """
        Create a Word document from chapters

        Args:
            title: Story title
            chapters: List of chapter dicts with 'chapter_number', 'title', 'content'
            author: Optional author name

        Returns:
            Document object
        """
        doc = Document()

        # Set up styles
        self._setup_styles(doc)

        # Separate Chapter 0 (intro) from other chapters
        intro_chapter = None
        regular_chapters = []
        for ch in chapters:
            if ch.get('chapter_number', 0) == 0:
                intro_chapter = ch
            else:
                regular_chapters.append(ch)

        # Add title page with intro content
        self._add_title_page(doc, title, author, len(regular_chapters), intro_chapter)

        # Add regular chapters (excluding Chapter 0)
        for chapter in sorted(regular_chapters, key=lambda x: x.get('chapter_number', 0)):
            self._add_chapter(doc, chapter)

        return doc

    def _setup_styles(self, doc: Document):
        """Set up custom styles for the document"""
        styles = doc.styles

        # Chapter title style
        if 'ChapterTitle' not in [s.name for s in styles]:
            chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.size = Pt(16)
            chapter_style.font.bold = True
            chapter_style.paragraph_format.space_before = Pt(24)
            chapter_style.paragraph_format.space_after = Pt(12)
            chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Content style
        if 'ChapterContent' not in [s.name for s in styles]:
            content_style = styles.add_style('ChapterContent', WD_STYLE_TYPE.PARAGRAPH)
            content_style.font.size = Pt(12)
            content_style.paragraph_format.space_after = Pt(8)
            content_style.paragraph_format.line_spacing = 1.5
            content_style.paragraph_format.first_line_indent = Inches(0.5)

    def _add_title_page(
        self,
        doc: Document,
        title: str,
        author: Optional[str],
        chapter_count: int,
        intro_chapter: Optional[dict] = None
    ):
        """Add title page to document"""
        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True

        # Add some space
        doc.add_paragraph()

        # Author if provided
        if author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"Tác giả: {author}")
            author_run.font.size = Pt(14)
            author_run.font.italic = True

        # Chapter count
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"Số chương: {chapter_count}")
        info_run.font.size = Pt(12)
        info_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add intro content if available
        if intro_chapter and intro_chapter.get('content'):
            doc.add_paragraph()  # Space before intro
            intro_content = intro_chapter.get('content', '')
            paragraphs = intro_content.split('\n\n') if intro_content else []
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph(style='ChapterContent')
                    self._add_runs_with_breaks(para, para_text)

        # Page break after title
        doc.add_page_break()

    def _add_runs_with_breaks(self, para, text: str):
        """Add text to paragraph, preserving single newlines as line breaks."""
        lines = text.split('\n')
        for i, line in enumerate(lines):
            para.add_run(line)
            if i < len(lines) - 1:
                para.add_run().add_break()

    def _add_chapter(self, doc: Document, chapter: dict):
        """Add a chapter to the document"""
        chapter_num = chapter.get('chapter_number', 0)
        chapter_title = chapter.get('title', f'Chương {chapter_num}')
        content = chapter.get('content', '')

        # Chapter title
        title_para = doc.add_paragraph(style='ChapterTitle')
        title_para.add_run(chapter_title)

        # Chapter content - split by paragraphs
        paragraphs = content.split('\n\n') if content else []
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph(style='ChapterContent')
                self._add_runs_with_breaks(para, para_text)

        # Add page break after each chapter (except last)
        doc.add_page_break()

    def export_story(
        self,
        story_id: str,
        title: str,
        chapters: List[dict],
        author: Optional[str] = None
    ) -> str:
        """
        Export story to Word document and save to file

        Args:
            story_id: Story ID for filename
            title: Story title
            chapters: List of chapter dicts
            author: Optional author name

        Returns:
            Path to saved file
        """
        # Create document
        doc = self.create_document(title, chapters, author)

        # Generate filename
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_title = safe_title[:50]  # Limit length
        filename = f"{safe_title}_{story_id[:8]}.docx"
        filepath = os.path.join(self.output_dir, filename)

        # Save document
        doc.save(filepath)
        logger.info(f"Exported story to: {filepath}")

        return filepath

    def get_file_size(self, filepath: str) -> int:
        """Get file size in bytes"""
        return os.path.getsize(filepath) if os.path.exists(filepath) else 0
